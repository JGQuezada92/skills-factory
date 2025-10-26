#!/usr/bin/env python3
"""
Investment Memo Document Exporter

Generates professional PDF and Microsoft Word documents from investment memos.
Outputs to the 'skills artifacts' folder for easy distribution.
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Optional


class MemoExporter:
    """Export investment memos to PDF and Word formats"""
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        Initialize exporter with output directory
        
        Args:
            output_dir: Directory to save files (default: skills artifacts folder)
        """
        if output_dir is None:
            # Default to skills artifacts folder
            self.output_dir = Path(
                "C:/Users/Jonathan Quezada/OneDrive - Phenom People, Inc/Desktop/Claude Skills/skills artifacts"
            )
        else:
            self.output_dir = Path(output_dir)
        
        # Create output directory if it doesn't exist
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export_to_markdown(self, content: str, company_name: str) -> str:
        """
        Save memo as Markdown file
        
        Args:
            content: The memo content in markdown format
            company_name: Name of the company for filename
            
        Returns:
            Path to saved file
        """
        try:
            # Clean company name for filename
            safe_name = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_name = safe_name.replace(' ', '_')
            
            # Create filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_name}_Investment_Memo_{timestamp}.md"
            filepath = self.output_dir / filename
            
            # Write content
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            print(f"[OK] Saved Markdown: {filepath}")
            return str(filepath)
            
        except Exception as e:
            print(f"[ERROR] Error saving Markdown: {e}")
            return None
    
    def export_to_pdf(self, content: str, company_name: str) -> str:
        """
        Convert memo to PDF format
        
        Args:
            content: The memo content in markdown format
            company_name: Name of the company for filename
            
        Returns:
            Path to saved PDF file
        """
        try:
            # Try to import PDF generation library
            try:
                from markdown2 import markdown
                from weasyprint import HTML, CSS
                has_pdf_support = True
            except ImportError:
                has_pdf_support = False
            
            if not has_pdf_support:
                # Fallback: Save instructions for manual conversion
                return self._create_pdf_instructions(content, company_name)
            
            # Clean company name for filename
            safe_name = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_name = safe_name.replace(' ', '_')
            
            # Create filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_name}_Investment_Memo_{timestamp}.pdf"
            filepath = self.output_dir / filename
            
            # Convert Markdown to HTML
            html_content = markdown(content, extras=['tables', 'fenced-code-blocks'])
            
            # Add professional styling - Modern presentation style with teal accents
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    @page {{
                        size: letter;
                        margin: 0.75in 1in;
                    }}
                    body {{
                        font-family: Arial, 'Helvetica Neue', sans-serif;
                        font-size: 11pt;
                        line-height: 1.6;
                        color: #2c3e50;
                        background: #ffffff;
                    }}
                    
                    /* Main Title - Presentation Style with Teal Header */
                    h1 {{
                        font-family: Arial, sans-serif;
                        font-size: 28pt;
                        font-weight: 300;
                        color: #2c9e91;
                        margin: 0 0 30px 0;
                        padding: 25px 30px;
                        background: linear-gradient(135deg, #2c9e91 0%, #45b8ac 100%);
                        color: white;
                        border-left: 8px solid #1a7a6f;
                        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                    }}
                    
                    /* Section Headers - Teal Accent */
                    h2 {{
                        font-family: Arial, sans-serif;
                        font-size: 16pt;
                        font-weight: 600;
                        color: #2c9e91;
                        margin-top: 35px;
                        margin-bottom: 15px;
                        padding-bottom: 8px;
                        border-bottom: 3px solid #2c9e91;
                    }}
                    
                    /* Sub-headers */
                    h3 {{
                        font-family: Arial, sans-serif;
                        font-size: 13pt;
                        font-weight: 600;
                        color: #34495e;
                        margin-top: 20px;
                        margin-bottom: 10px;
                    }}
                    
                    /* Key Metrics Boxes - Similar to "WHY NOW" style */
                    .key-metric {{
                        background: #f8fffe;
                        border-left: 4px solid #2c9e91;
                        padding: 15px 20px;
                        margin: 15px 0;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.08);
                    }}
                    
                    /* Executive Summary Box */
                    .exec-summary {{
                        background: linear-gradient(135deg, #f0f9f8 0%, #ffffff 100%);
                        border: 2px solid #2c9e91;
                        border-radius: 8px;
                        padding: 25px;
                        margin: 25px 0;
                        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
                    }}
                    
                    /* Professional Tables */
                    table {{
                        width: 100%;
                        border-collapse: collapse;
                        margin: 25px 0;
                        font-size: 10pt;
                        background: white;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.08);
                    }}
                    th, td {{
                        border: 1px solid #e0e0e0;
                        padding: 12px 15px;
                        text-align: left;
                    }}
                    th {{
                        background: linear-gradient(135deg, #2c9e91 0%, #45b8ac 100%);
                        color: white;
                        font-weight: 600;
                        text-transform: uppercase;
                        font-size: 9pt;
                        letter-spacing: 0.5px;
                    }}
                    tr:nth-child(even) {{
                        background-color: #f8fffe;
                    }}
                    tr:hover {{
                        background-color: #f0f9f8;
                    }}
                    
                    /* Bullet Points - Clean Style */
                    ul {{
                        list-style-type: none;
                        padding-left: 0;
                        margin: 15px 0;
                    }}
                    ul li {{
                        padding-left: 30px;
                        margin: 10px 0;
                        position: relative;
                        line-height: 1.6;
                    }}
                    ul li:before {{
                        content: "▸";
                        color: #2c9e91;
                        font-weight: bold;
                        position: absolute;
                        left: 10px;
                        font-size: 14pt;
                    }}
                    
                    /* Numbered Lists */
                    ol {{
                        counter-reset: item;
                        padding-left: 0;
                        margin: 15px 0;
                    }}
                    ol li {{
                        counter-increment: item;
                        padding-left: 40px;
                        margin: 12px 0;
                        position: relative;
                    }}
                    ol li:before {{
                        content: counter(item);
                        background: #2c9e91;
                        color: white;
                        font-weight: bold;
                        font-size: 10pt;
                        border-radius: 50%;
                        width: 24px;
                        height: 24px;
                        display: inline-flex;
                        align-items: center;
                        justify-content: center;
                        position: absolute;
                        left: 0;
                    }}
                    
                    /* Emphasis */
                    strong {{
                        color: #2c3e50;
                        font-weight: 600;
                    }}
                    
                    em {{
                        color: #2c9e91;
                        font-style: normal;
                        font-weight: 600;
                    }}
                    
                    /* Blockquotes - Styled like call-out boxes */
                    blockquote {{
                        border-left: 6px solid #2c9e91;
                        background: #f8fffe;
                        margin: 20px 0;
                        padding: 20px 25px;
                        font-style: italic;
                        color: #34495e;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.08);
                    }}
                    
                    /* Code blocks */
                    code {{
                        background-color: #ecf0f1;
                        padding: 2px 8px;
                        font-family: 'Courier New', monospace;
                        font-size: 10pt;
                        color: #e74c3c;
                        border-radius: 3px;
                    }}
                    
                    /* Horizontal Rules */
                    hr {{
                        border: none;
                        height: 2px;
                        background: linear-gradient(90deg, #2c9e91 0%, transparent 100%);
                        margin: 30px 0;
                    }}
                    
                    /* Page breaks for printing */
                    .page-break {{
                        page-break-after: always;
                    }}
                    
                    /* Footer styling */
                    .footer {{
                        margin-top: 40px;
                        padding-top: 20px;
                        border-top: 2px solid #ecf0f1;
                        font-size: 9pt;
                        color: #7f8c8d;
                        text-align: center;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Generate PDF
            HTML(string=styled_html).write_pdf(filepath)
            
            print(f"[OK] Saved PDF: {filepath}")
            return str(filepath)
            
        except Exception as e:
            print(f"[ERROR] Error generating PDF: {e}")
            return self._create_pdf_instructions(content, company_name)
    
    def export_to_word(self, content: str, company_name: str) -> str:
        """
        Convert memo to Microsoft Word format
        
        Args:
            content: The memo content in markdown format
            company_name: Name of the company for filename
            
        Returns:
            Path to saved Word file
        """
        try:
            # Try to import Word generation library
            try:
                from docx import Document
                from docx.shared import Pt, Inches, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                has_word_support = True
            except ImportError:
                has_word_support = False
            
            if not has_word_support:
                # Fallback: Save instructions for manual conversion
                return self._create_word_instructions(content, company_name)
            
            # Clean company name for filename
            safe_name = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_name = safe_name.replace(' ', '_')
            
            # Create filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_name}_Investment_Memo_{timestamp}.docx"
            filepath = self.output_dir / filename
            
            # Create Word document with professional styling
            doc = Document()
            
            # Set default font - Arial throughout
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(10)
            style.paragraph_format.line_spacing = 1.15
            
            # Customize heading styles with teal color
            # Heading 1 - Main Title
            heading1_style = doc.styles['Heading 1']
            heading1_style.font.name = 'Arial'
            heading1_style.font.size = Pt(28)
            heading1_style.font.color.rgb = RGBColor(44, 158, 145)  # Teal
            heading1_style.paragraph_format.space_before = Pt(0)
            heading1_style.paragraph_format.space_after = Pt(20)
            
            # Heading 2 - Section Headers
            heading2_style = doc.styles['Heading 2']
            heading2_style.font.name = 'Arial'
            heading2_style.font.size = Pt(16)
            heading2_style.font.bold = True
            heading2_style.font.color.rgb = RGBColor(44, 158, 145)  # Teal
            heading2_style.paragraph_format.space_before = Pt(18)
            heading2_style.paragraph_format.space_after = Pt(10)
            
            # Heading 3 - Sub-headers
            heading3_style = doc.styles['Heading 3']
            heading3_style.font.name = 'Arial'
            heading3_style.font.size = Pt(13)
            heading3_style.font.bold = True
            heading3_style.font.color.rgb = RGBColor(52, 73, 94)  # Dark gray
            heading3_style.paragraph_format.space_before = Pt(14)
            heading3_style.paragraph_format.space_after = Pt(8)
            
            # Parse markdown and add to document
            self._markdown_to_docx(content, doc)
            
            # Save document
            doc.save(filepath)
            
            print(f"[OK] Saved Word Document: {filepath}")
            return str(filepath)
            
        except Exception as e:
            print(f"[ERROR] Error generating Word document: {e}")
            return self._create_word_instructions(content, company_name)
    
    def _markdown_to_docx(self, content: str, doc):
        """Convert markdown content to Word document"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            
            # Numbered lists
            elif line[0].isdigit() and line[1:3] == '. ':
                doc.add_paragraph(line[3:], style='List Number')
            
            # Horizontal rule
            elif line == '---':
                doc.add_paragraph('_' * 50)
            
            # Regular paragraph
            else:
                doc.add_paragraph(line)
    
    def _create_pdf_instructions(self, content: str, company_name: str) -> str:
        """Create instructions file for manual PDF conversion"""
        # Save markdown first
        md_path = self.export_to_markdown(content, company_name)
        
        instructions = f"""
PDF Export Instructions
========================

The required PDF libraries are not installed. To generate PDF:

Option 1: Install libraries and re-run
---------------------------------------
pip install markdown2 weasyprint

Option 2: Manual conversion
----------------------------
1. Open the Markdown file: {md_path}
2. Use an online converter: https://www.markdowntopdf.com/
3. Or use Pandoc: pandoc input.md -o output.pdf

Option 3: Print to PDF from Claude
-----------------------------------
1. View the memo in Claude
2. Use your browser's Print function (Ctrl+P)
3. Select "Save as PDF"
4. Save to: {self.output_dir}
"""
        
        instructions_file = self.output_dir / "PDF_EXPORT_INSTRUCTIONS.txt"
        with open(instructions_file, 'w', encoding='utf-8') as f:
            f.write(instructions)
        
        print(f"[WARN] PDF libraries not available. Instructions saved: {instructions_file}")
        return str(instructions_file)
    
    def _create_word_instructions(self, content: str, company_name: str) -> str:
        """Create instructions file for manual Word conversion"""
        # Save markdown first
        md_path = self.export_to_markdown(content, company_name)
        
        instructions = f"""
Word Export Instructions
=========================

The required Word libraries are not installed. To generate Word document:

Option 1: Install libraries and re-run
---------------------------------------
pip install python-docx

Option 2: Manual conversion
----------------------------
1. Open the Markdown file: {md_path}
2. Copy the content
3. Paste into Microsoft Word
4. Apply formatting as needed
5. Save to: {self.output_dir}

Option 3: Use Pandoc
---------------------
pandoc input.md -o output.docx
"""
        
        instructions_file = self.output_dir / "WORD_EXPORT_INSTRUCTIONS.txt"
        with open(instructions_file, 'w', encoding='utf-8') as f:
            f.write(instructions)
        
        print(f"[WARN] Word libraries not available. Instructions saved: {instructions_file}")
        return str(instructions_file)
    
    def export_all(self, content: str, company_name: str) -> dict:
        """
        Export memo to all formats
        
        Args:
            content: The memo content in markdown format
            company_name: Name of the company
            
        Returns:
            Dictionary with paths to all exported files
        """
        results = {
            'markdown': self.export_to_markdown(content, company_name),
            'pdf': self.export_to_pdf(content, company_name),
            'word': self.export_to_word(content, company_name)
        }
        
        print("\n" + "="*70)
        print("Export Complete! Files saved to:")
        print(str(self.output_dir))
        print("="*70 + "\n")
        
        return results


# Convenience function for direct use
def export_memo(content: str, company_name: str, formats: list = ['markdown', 'pdf', 'word']) -> dict:
    """
    Export investment memo to specified formats
    
    Args:
        content: The memo content in markdown format
        company_name: Name of the company for filename
        formats: List of formats to export ('markdown', 'pdf', 'word')
        
    Returns:
        Dictionary with paths to exported files
    """
    exporter = MemoExporter()
    results = {}
    
    if 'markdown' in formats:
        results['markdown'] = exporter.export_to_markdown(content, company_name)
    if 'pdf' in formats:
        results['pdf'] = exporter.export_to_pdf(content, company_name)
    if 'word' in formats:
        results['word'] = exporter.export_to_word(content, company_name)
    
    return results


if __name__ == "__main__":
    # Example usage
    sample_memo = """
# EXAMPLE COMPANY - INVESTMENT MEMORANDUM

## INVESTMENT RECOMMENDATION: BUY

**Price Target:** $100  
**Current Price:** $50  
**Upside:** 100%

## EXECUTIVE SUMMARY

This is a sample investment memo demonstrating the export functionality.

### Key Points:
- Strong revenue growth
- Expanding margins
- Large TAM opportunity

## COMPANY OVERVIEW

Example Company is a leading provider of innovative solutions...
"""
    
    print("Testing memo export functionality...")
    exporter = MemoExporter()
    results = exporter.export_all(sample_memo, "Example Company")
    
    print("\nExport test complete!")
    print(f"Files created: {len(results)}")

